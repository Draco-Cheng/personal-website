"""
File parsing utilities for extracting text from various document formats.
Supports: PDF, DOCX, XLSX, Markdown, and plain text files.
"""

from typing import Optional
import io

def parse_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF file.
    First tries PyPDF2, falls back to pdfplumber for complex PDFs.

    Args:
        file_content: Raw bytes of the PDF file

    Returns:
        Extracted text content

    Raises:
        Exception: If PDF parsing fails
    """
    try:
        # Try PyPDF2 first (faster, lighter)
        from PyPDF2 import PdfReader

        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)

        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        result = "\n\n".join(text_parts)

        # If extraction yielded very little text, try pdfplumber
        if len(result.strip()) < 50 and len(file_content) > 1000:
            return parse_pdf_with_pdfplumber(file_content)

        return result

    except Exception as e:
        # Fallback to pdfplumber for complex PDFs
        try:
            return parse_pdf_with_pdfplumber(file_content)
        except Exception as fallback_error:
            raise Exception(f"PDF parsing failed: {str(e)}, Fallback also failed: {str(fallback_error)}")


def parse_pdf_with_pdfplumber(file_content: bytes) -> str:
    """
    Extract text from PDF using pdfplumber (better for complex layouts).

    Args:
        file_content: Raw bytes of the PDF file

    Returns:
        Extracted text content
    """
    import pdfplumber

    pdf_file = io.BytesIO(file_content)
    text_parts = []

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

    return "\n\n".join(text_parts)


def parse_docx(file_content: bytes) -> str:
    """
    Extract text from DOCX (Word) file.

    Args:
        file_content: Raw bytes of the DOCX file

    Returns:
        Extracted text content
    """
    from docx import Document

    docx_file = io.BytesIO(file_content)
    doc = Document(docx_file)

    text_parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))

    return "\n\n".join(text_parts)


def parse_xlsx(file_content: bytes) -> str:
    """
    Extract text from XLSX (Excel) file.
    Converts each sheet to text format.

    Args:
        file_content: Raw bytes of the XLSX file

    Returns:
        Extracted text content with sheet names
    """
    from openpyxl import load_workbook

    xlsx_file = io.BytesIO(file_content)
    workbook = load_workbook(xlsx_file, data_only=True)

    text_parts = []

    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        text_parts.append(f"=== Sheet: {sheet_name} ===")

        for row in sheet.iter_rows(values_only=True):
            # Filter out None values and convert to strings
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                text_parts.append(" | ".join(row_values))

    return "\n\n".join(text_parts)


def parse_markdown(file_content: bytes) -> str:
    """
    Extract text from Markdown file.
    Simply decodes as UTF-8 text.

    Args:
        file_content: Raw bytes of the Markdown file

    Returns:
        Decoded text content
    """
    return file_content.decode("utf-8", errors="ignore")


def parse_text(file_content: bytes) -> str:
    """
    Extract text from plain text file.
    Tries UTF-8, falls back to other encodings.

    Args:
        file_content: Raw bytes of the text file

    Returns:
        Decoded text content
    """
    # Try UTF-8 first
    try:
        return file_content.decode("utf-8")
    except UnicodeDecodeError:
        # Try other common encodings
        for encoding in ["latin-1", "cp1252", "iso-8859-1"]:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue

    # Last resort: decode with errors ignored
    return file_content.decode("utf-8", errors="ignore")


def parse_file(file_content: bytes, filename: str) -> str:
    """
    Parse a file based on its extension.

    Args:
        file_content: Raw bytes of the file
        filename: Original filename (used to determine file type)

    Returns:
        Extracted text content

    Raises:
        ValueError: If file type is not supported
        Exception: If parsing fails
    """
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return parse_pdf(file_content)
    elif filename_lower.endswith(".docx"):
        return parse_docx(file_content)
    elif filename_lower.endswith(".xlsx"):
        return parse_xlsx(file_content)
    elif filename_lower.endswith(".md") or filename_lower.endswith(".markdown"):
        return parse_markdown(file_content)
    elif filename_lower.endswith(".txt"):
        return parse_text(file_content)
    else:
        raise ValueError(f"Unsupported file type: {filename}")


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing excessive whitespace and special characters.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    # Remove multiple consecutive newlines
    lines = text.split("\n")
    cleaned_lines = []

    for line in lines:
        # Strip whitespace
        line = line.strip()
        # Skip empty lines or lines with only special characters
        if line and not all(c in " \t\n\r" for c in line):
            cleaned_lines.append(line)

    # Join with single newlines, but preserve paragraph breaks
    result = []
    for i, line in enumerate(cleaned_lines):
        result.append(line)
        # Add extra newline for paragraph breaks (heuristic)
        if i < len(cleaned_lines) - 1:
            next_line = cleaned_lines[i + 1]
            # If current line ends with punctuation and next starts with capital, likely paragraph break
            if line and line[-1] in ".!?" and next_line and next_line[0].isupper():
                result.append("")

    return "\n".join(result)
